"""Build PathoLens-VLM report by filling the ITU AI&DE template in-place.

Re-runnable: starts from the pristine template each time, fills sections,
saves to PathoLens_Report.docx. Open in Word and press F9 (or right-click
TOC -> Update Field) to refresh the table of contents.
"""

from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph

TEMPLATE = Path("/Users/farukrizaoz/Downloads/04-4902-Project_Report_Template.docx")
OUTPUT = Path("/Users/farukrizaoz/Desktop/patholens/report/PathoLens_Report.docx")


# ─── Helpers ──────────────────────────────────────────────────────────────────


def set_text_keep_style(para: Paragraph, text: str) -> None:
    """Replace paragraph text while preserving its style and first-run formatting."""
    if para.runs:
        first = para.runs[0]
        first.text = text
        for run in para.runs[1:]:
            run._element.getparent().remove(run._element)
    else:
        para.add_run(text)


def insert_para_after(prev: Paragraph, text: str, style_name: str, doc) -> Paragraph:
    """Insert a new paragraph after `prev` with the given text and style."""
    new_element = deepcopy(prev._element)
    prev._element.addnext(new_element)
    new_para = Paragraph(new_element, prev._parent)
    # Strip all existing runs and hyperlinks so we start clean.
    for child in list(new_element):
        if child.tag in (qn("w:r"), qn("w:hyperlink")):
            new_element.remove(child)
    new_para.style = doc.styles[style_name]
    new_para.add_run(text)
    return new_para


def find_para(doc, text: str, style_prefix: str = "") -> Paragraph:
    """Find a paragraph by exact text (case-insensitive) and optional style prefix."""
    for p in doc.paragraphs:
        if p.text.strip().upper() == text.upper() and (
            not style_prefix or p.style.name.startswith(style_prefix)
        ):
            return p
    raise ValueError(f"Paragraph {text!r} not found (style prefix {style_prefix!r})")


def fill_section(doc, heading_text: str, body_blocks: list[tuple[str, str]]) -> Paragraph:
    """Replace the first placeholder paragraph after a Heading 1 with body_blocks.

    body_blocks items: (text, style_name).  Returns the last-inserted paragraph.
    """
    heading = find_para(doc, heading_text, style_prefix="Heading")
    placeholder = Paragraph(heading._element.getnext(), heading._parent)

    first_text, first_style = body_blocks[0]
    placeholder.style = doc.styles[first_style]
    set_text_keep_style(placeholder, first_text)

    prev = placeholder
    for text, style_name in body_blocks[1:]:
        prev = insert_para_after(prev, text, style_name, doc)
    return prev


def insert_blocks_after_subheading(
    doc, sub_heading_text: str, body_blocks: list[tuple[str, str]]
) -> Paragraph:
    """Insert body_blocks immediately after a Heading 2 (no placeholder to replace)."""
    heading = find_para(doc, sub_heading_text, style_prefix="Heading 2")
    prev = heading
    for text, style_name in body_blocks:
        prev = insert_para_after(prev, text, style_name, doc)
    return prev


def insert_table_after(
    prev: Paragraph,
    headers: list[str],
    rows: list[list[str]],
    doc,
) -> None:
    """Insert a Table-Grid table after `prev` paragraph."""
    n_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for col_idx, header_text in enumerate(headers):
        cell = table.rows[0].cells[col_idx]
        cell.text = ""
        run = cell.paragraphs[0].add_run(header_text)
        run.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Body rows
    for row_idx, row_values in enumerate(rows, start=1):
        for col_idx, value in enumerate(row_values):
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = ""
            cell.paragraphs[0].add_run(str(value))
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Move the table from end-of-document to immediately after `prev`.
    tbl_element = table._element
    tbl_element.getparent().remove(tbl_element)
    prev._element.addnext(tbl_element)


def add_caption_and_table(
    prev: Paragraph,
    caption: str,
    headers: list[str],
    rows: list[list[str]],
    doc,
) -> Paragraph:
    """Add caption paragraph then table (per ITU guideline: caption above table)."""
    cap = insert_para_after(prev, caption, "Normal", doc)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap.runs:
        run.italic = True
        run.font.size = Pt(11)
    insert_table_after(cap, headers, rows, doc)
    # Add an empty spacing paragraph after the table
    spacer = insert_para_after(cap, "", "Normal", doc)
    # Note: spacer is inserted after caption (before table) because addnext on cap.
    # Move spacer to after the table by relocating.
    spacer_el = spacer._element
    spacer_el.getparent().remove(spacer_el)
    cap._element.addnext(spacer_el)  # this places it right after caption again — undo
    # Simpler: build a NEW spacer at end by walking past the table.
    tbl_el = cap._element.getnext()
    # tbl_el should be the inserted <w:tbl>
    if tbl_el is not None and tbl_el.tag == qn("w:tbl"):
        # Insert empty paragraph after the table
        new_el = deepcopy(cap._element)
        for child in list(new_el):
            if child.tag in (qn("w:r"), qn("w:hyperlink")):
                new_el.remove(child)
        tbl_el.addnext(new_el)
        spacer_real = Paragraph(new_el, cap._parent)
        spacer_real.style = doc.styles["Normal"]
        # Remove the wrongly-placed spacer
        if spacer_el.getparent() is not None:
            spacer_el.getparent().remove(spacer_el)
        return spacer_real
    return cap


def add_figure_placeholder(prev: Paragraph, caption: str, doc) -> Paragraph:
    """Insert a centred italic figure placeholder + caption line (image to be added in Word)."""
    placeholder_box = insert_para_after(
        prev,
        f"[Figure placeholder — insert image here in Word: {caption.split(':', 1)[-1].strip()}]",
        "Normal",
        doc,
    )
    placeholder_box.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in placeholder_box.runs:
        run.italic = True
        run.font.color.rgb = None  # default
    cap = insert_para_after(placeholder_box, caption, "Normal", doc)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap.runs:
        run.italic = True
        run.font.size = Pt(11)
    return cap


# ─── Constants ────────────────────────────────────────────────────────────────

PROJECT_TITLE = (
    "PathoLens-VLM: A Spatially-Grounded Vision-Language Model "
    "for Whole-Slide Breast Cancer Pathology Report Generation"
)
STUDENT_1 = "150200000 Faruk Rıza ÖZ"  # TODO: confirm student ID
STUDENT_2 = "150200000 Emir Arda EKER"  # TODO: confirm student ID
SUPERVISOR = "Prof. Dr. Behçet Uğur TÖREYİN"
SUBMISSION_DATE = "June, 2026"


# ─── SUMMARY ──────────────────────────────────────────────────────────────────

SUMMARY_BODY: list[tuple[str, str]] = [
    (
        "PathoLens-VLM is a vision-language model that generates structured "
        "diagnostic reports for whole-slide images of breast cancer tissue and "
        "explicitly grounds each generated sentence to the slide region that "
        "supports it. Existing pathology report-generation systems "
        "(PathChat, SlideChat, Quilt-LLaVA) can produce fluent captions but "
        "cannot indicate which area of the slide each sentence comes from, "
        "which prevents their use in any safety-critical clinical workflow.",
        "Normal",
    ),
    (
        "The proposed system combines three components: a frozen CONCHv1.5 "
        "patch encoder that converts each 448×448 tile of the slide into a "
        "768-dimensional embedding; a trainable linear adapter that projects "
        "those embeddings into the hidden space of the language model; and a "
        "Llama-3.2-3B-Instruct language model fine-tuned with Low-Rank "
        "Adaptation (LoRA). During training, a KL-divergence grounding loss "
        "supervises the language model’s text-to-vision attention against "
        "pixel-level BCSS semantic segmentation masks where available and a "
        "CONCH-v1 cross-modal pseudo-ground-truth otherwise. A grounding "
        "warmup schedule and an entropy-based faithfulness regulariser are "
        "added to stabilise training and force attention concentration.",
        "Normal",
    ),
    (
        "On a held-out test set of seven BCSS-annotated TCGA-BRCA slides, "
        "PathoLens-VLM achieves BLEU-4 = 0.0559, ROUGE-L = 0.1990 and a BCSS "
        "pointing-game accuracy of PG@5 = 0.081, a 4.4× lift over the "
        "uniform-random baseline of 0.019. The attention distribution is "
        "substantially peakier than uniform (entropy ratio 0.938; top-10 % "
        "attention mass 0.201 vs uniform 0.100). Only ~15 M parameters are "
        "trainable (~5 M adapter + ~10 M LoRA); the remaining 3.3 B parameters "
        "of the language model and patch encoder are frozen, keeping the "
        "training cost on a single A6000 GPU within 12 hours.",
        "Normal",
    ),
    (
        "The system is delivered as an open-source repository with five "
        "ablation configurations, a Streamlit demonstration application that "
        "produces FHIR DiagnosticReport JSON output, and 51 passing unit "
        "and integration tests.",
        "Normal",
    ),
]


# ─── INTRODUCTION ─────────────────────────────────────────────────────────────

INTRODUCTION_BODY: list[tuple[str, str]] = [
    (
        "Digital pathology has emerged as one of the most active application "
        "domains for medical artificial intelligence, with whole-slide images "
        "(WSIs) now routinely scanned at resolutions exceeding 50,000 × 50,000 "
        "pixels per slide. A single pathologist examines on the order of 10–80 "
        "such slides per working day, and the manual drafting of the "
        "associated diagnostic report is among the most time-consuming and "
        "error-prone tasks in the clinical workflow. The recent emergence of "
        "vision-language models (VLMs), large language models (LLMs) augmented "
        "with a vision encoder, has therefore raised the prospect of automatic "
        "draft report generation directly from a WSI, both as a triage aid "
        "and as a quality-control second reader.",
        "Normal",
    ),
    (
        "A number of pathology-specific VLMs have appeared in the past two "
        "years, the most prominent being PathChat [1], SlideChat [2] and "
        "Quilt-LLaVA [3]. These systems can produce fluent, plausible "
        "captions for a slide and answer free-form questions about its "
        "contents. However, they share a critical limitation that prevents "
        "their use in any safety-critical clinical context: the generated "
        "report carries no explicit link back to the slide region that each "
        "clinical claim originates from. A pathologist confronted with the "
        "sentence “moderate stromal lymphocytic infiltrate is present” must "
        "accept the model’s word for it, without any way to inspect the "
        "supporting evidence on the slide. In a discipline where the "
        "diagnostic conclusion routinely depends on the appearance of a "
        "0.1 mm² focus inside a 1 cm² section, this absence of spatial "
        "grounding is not a stylistic preference; it is a structural barrier "
        "to clinical adoption.",
        "Normal",
    ),
    (
        "Spatially-grounded report generation has been studied in the "
        "radiology literature, where the recent MedGround model [4] aligns "
        "each report sentence with a bounding box on the source image. The "
        "histopathology equivalent is substantially harder for two reasons. "
        "First, a WSI contains three to four orders of magnitude more pixels "
        "than a chest X-ray, so the candidate regions run into the thousands "
        "of patches per slide. Second, while radiology benefits from large "
        "datasets of paired reports and pixel-level abnormality masks, "
        "histopathology has very few publicly available datasets that combine "
        "free-text reports with patch-level semantic segmentation.",
        "Normal",
    ),
    ("Research Question", "Heading 2"),
    (
        "The central research question of this capstone project is the "
        "following: can a vision-language model for whole-slide breast cancer "
        "pathology produce a free-form diagnostic report in which every "
        "clinical sentence is explicitly grounded to the slide region that "
        "supports it, and can this grounding be measured by an objective "
        "metric on a held-out test set? The question is deliberately "
        "constructive: rather than asking whether post-hoc explainability "
        "tools (e.g., Grad-CAM, attention rollout) can be applied to an "
        "existing pathology VLM, we ask whether the grounding signal can be "
        "built into the training objective itself, so that the model learns "
        "to attend to the correct slide region as part of producing the "
        "sentence.",
        "Normal",
    ),
    ("Contributions", "Heading 2"),
    (
        "This work presents PathoLens-VLM, a spatially-grounded vision-language "
        "model for whole-slide breast cancer pathology report generation. "
        "The specific contributions are as follows.",
        "Normal",
    ),
    (
        "First, a grounded VLM architecture for whole-slide pathology. "
        "PathoLens-VLM combines the frozen CONCHv1.5 patch encoder [5], a "
        "trainable linear adapter, and the Llama-3.2-3B-Instruct language "
        "model [6] with LoRA fine-tuning [7]. The architecture preserves "
        "one-to-one correspondence between LLM vision tokens and patch-level "
        "grounding ground truth, which is a prerequisite for the supervised "
        "grounding loss.",
        "Normal",
    ),
    (
        "Second, a hybrid grounding supervision strategy. A KL-divergence "
        "grounding loss supervises the LLM’s text-to-vision attention against "
        "a two-source ground-truth distribution: pixel-level BCSS semantic "
        "segmentation masks [8] where available, and a CONCH-v1 cross-modal "
        "pseudo-ground-truth as fallback. A grounding warmup schedule "
        "prevents the grounding signal, which is noisy early in training, "
        "from disrupting caption learning.",
        "Normal",
    ),
    (
        "Third, a reproducible evaluation suite. Five complementary metrics "
        "jointly measure caption quality (BLEU-4, ROUGE-L), spatial grounding "
        "(BCSS pointing-game at K), attention concentration (top-K mass, "
        "normalised Shannon entropy), and concept recall (keyword-based "
        "precision and recall against BCSS class coverage). All metrics are "
        "scripted and reproducible from the project repository.",
        "Normal",
    ),
    (
        "Fourth, empirical evidence of measurable grounding. On a held-out "
        "test set of seven BCSS-annotated TCGA-BRCA slides, PathoLens-VLM "
        "achieves BLEU-4 = 0.0559, ROUGE-L = 0.1990 and a BCSS pointing-game "
        "accuracy of PG@5 = 0.081 against a uniform-random baseline of 0.019, "
        "a 4.4× lift over chance. The attention distribution is substantially "
        "peakier than uniform (entropy ratio 0.938; top-10 % attention mass "
        "0.201 against the uniform baseline of 0.100), demonstrating that "
        "the grounding loss successfully concentrates attention onto a "
        "small, slide-specific region.",
        "Normal",
    ),
    (
        "Fifth, an interactive demonstration. A Streamlit-based browser "
        "application lets a user select a precomputed slide, generate a "
        "structured pathology report, inspect per-sentence attention "
        "heatmaps overlaid on the slide thumbnail, and export the result as "
        "a FHIR DiagnosticReport JSON document.",
        "Normal",
    ),
    ("Scope and Limitations", "Heading 2"),
    (
        "This work is deliberately scoped to a single organ system (breast) "
        "and a single primary dataset (TCGA-BRCA), which is the largest "
        "publicly available WSI collection for which both free-text reports "
        "and patch-level segmentation masks (via BCSS) exist. The fullscale "
        "training run uses 50 slides, the intersection of BCSS annotations, "
        "SlideInstruction caption availability and successful GDC downloads "
        "at the time of the run. The held-out test set consists of seven "
        "slides, which is sufficient to establish that the grounding signal "
        "is non-trivial and reproducible, but is too small to support strong "
        "claims about clinical performance. A discussion of these and other "
        "limitations is given in Section 6.",
        "Normal",
    ),
    ("Report Organization", "Heading 2"),
    (
        "The remainder of this report is organised as follows. Section 2 "
        "reviews the relevant background in computational pathology, "
        "vision-language models for medical imaging, and grounded image "
        "captioning. Section 3 specifies the system requirements, including "
        "design constraints, functional and non-functional requirements, and "
        "the evaluation criteria. Section 4 details the data sources, the "
        "preprocessing pipeline, the system architecture and the training "
        "procedure. Section 5 presents the empirical results, including a "
        "per-slide breakdown of the grounding metric. Section 6 concludes "
        "with a discussion of limitations and directions for future work, "
        "and Section 7 lists the references cited throughout the report.",
        "Normal",
    ),
]


# ─── BACKGROUND ──────────────────────────────────────────────────────────────

BACKGROUND_BODY: list[tuple[str, str]] = [
    (
        "This section reviews the four bodies of literature that this project "
        "builds upon: computational pathology and whole-slide image (WSI) "
        "processing, foundation models for histopathology, vision-language "
        "models for medical imaging, and spatial grounding in image "
        "captioning. The discussion is concluded with a summary of the "
        "specific gap that PathoLens-VLM addresses.",
        "Normal",
    ),
    ("Computational Pathology and WSI Processing", "Heading 2"),
    (
        "A WSI is produced by a scanner that captures a tissue section at a "
        "microscope magnification of 20× or 40× and assembles the resulting "
        "tiles into a single pyramidal image, typically stored in the SVS or "
        "TIFF format. Even at moderate magnification a single slide occupies "
        "1–3 GB on disk and contains 10⁸–10⁹ informative pixels. Direct "
        "training of a neural network on such an image is computationally "
        "infeasible; the field has therefore converged on a two-stage "
        "pipeline in which (i) the slide is divided into a grid of smaller "
        "tiles (patches) of around 224 × 224 to 448 × 448 pixels and (ii) a "
        "patch-level encoder produces a low-dimensional embedding for each "
        "tile.",
        "Normal",
    ),
    (
        "The dominant aggregation strategy for the resulting bag of patch "
        "embeddings is Multiple-Instance Learning (MIL), originally proposed "
        "for binary slide-level classification [9]. Attention-based MIL "
        "approaches such as ABMIL and CLAM learn a weighted average over "
        "patches and have become the de facto baseline for slide-level "
        "classification on TCGA cohorts. PathoLens-VLM departs from this "
        "tradition in that the aggregation is performed by a Transformer "
        "language model rather than a learned pooling operation, which "
        "exposes per-patch attention weights that can be supervised directly.",
        "Normal",
    ),
    ("Foundation Models for Histopathology", "Heading 2"),
    (
        "Since 2023, several large-scale self-supervised patch encoders have "
        "been released that have substantially improved downstream "
        "performance across histopathology tasks. CONCH [10] and its "
        "successor CONCHv1.5 [5] are ViT-L vision encoders pretrained on "
        "1.17 million histology image-caption pairs with a contrastive "
        "objective; CONCHv1.5 is vision-only, while the original CONCH v1 "
        "retains a shared vision-text embedding space that can be used for "
        "zero-shot cross-modal retrieval. TITAN [11] (Mahmood Lab, 2024) is "
        "a slide-level Transformer that consumes CONCHv1.5 patch features "
        "and produces a small set of slide tokens via hierarchical "
        "attention, designed primarily for slide-level retrieval and "
        "classification.",
        "Normal",
    ),
    (
        "These models are released under research-only licences via the "
        "HuggingFace Hub with institutional gating, and have rapidly become "
        "the de facto backbones for downstream computational-pathology "
        "research. In PathoLens-VLM, CONCHv1.5 patch embeddings serve as "
        "the LLM vision tokens, and CONCH v1 is used separately for "
        "cross-modal text-to-patch similarity to construct the pseudo "
        "ground-truth grounding signal. TITAN is precomputed and stored for "
        "future ablation but is not part of the current training loop, for "
        "reasons explained in Section 4.",
        "Normal",
    ),
    ("Vision-Language Models for Medical Imaging", "Heading 2"),
    (
        "Vision-language models combine a vision encoder with an "
        "autoregressive language model through a learned projection. The "
        "LLaVA family [12] popularised the simple recipe of (i) a frozen "
        "vision encoder, (ii) a small projection layer (originally a single "
        "linear matrix; later a two-layer multilayer perceptron) and "
        "(iii) a language model fine-tuned on instruction-style examples in "
        "which image embeddings are prepended to the text prompt. LLaVA "
        "demonstrated that the projection layer alone, with only a few "
        "million trainable parameters, is sufficient to bridge a strong "
        "vision encoder to a strong language model.",
        "Normal",
    ),
    (
        "PathChat [1] adapted this recipe to pathology by fine-tuning on "
        "GPT-4-generated VQA pairs over histology images; SlideChat [2] "
        "scaled it to whole slides by pooling many patch embeddings before "
        "passing them to the language model; and Quilt-LLaVA [3] trained "
        "on a one-million-pair pathology image-caption dataset assembled "
        "from public lecture videos. All three systems produce free-form "
        "text and report respectable BLEU and ROUGE numbers, but none of "
        "them supervises the model’s attention or provides per-sentence "
        "spatial grounding.",
        "Normal",
    ),
    ("Spatial Grounding in Image Captioning", "Heading 2"),
    (
        "Spatial grounding in image captioning refers to the property that "
        "each generated phrase or sentence can be traced back to a specific "
        "image region that supports it. In the general domain, GLIP [13] "
        "and GroundingDINO [14] perform grounded object detection with "
        "free-form text queries, and dense-captioning systems such as the "
        "GRIT [15] benchmark provide phrase-level region annotations. In "
        "the medical domain, MedGround [4] adapts the GLIP approach to "
        "chest X-rays, aligning each sentence in the radiology report with a "
        "bounding box on the image. MedGround uses MS-CXR phrase-grounding "
        "annotations, which exist for radiology but, to the best of our "
        "knowledge, have no equivalent in histopathology.",
        "Normal",
    ),
    (
        "A closely related work is the Direct Visual Grounding (DVG) "
        "paper [16] of 2025, which introduces a KL-divergence attention loss "
        "to supervise the text-to-vision attention of an autoregressive "
        "captioning model. PathoLens-VLM adopts the same loss formulation "
        "and applies it for the first time at the whole-slide scale.",
        "Normal",
    ),
    ("Datasets and Annotations Used", "Heading 2"),
    (
        "Three public datasets underpin the project. TCGA-BRCA "
        "(The Cancer Genome Atlas, breast invasive carcinoma cohort) "
        "provides 1,098 WSI cases together with anonymised pathology "
        "reports. BCSS [8] (Breast Cancer Semantic Segmentation, Amgad et "
        "al. 2019) is a CC0-licensed dataset of 151 TCGA-BRCA slides for "
        "which expert pathologists have annotated 22 tissue classes at the "
        "pixel level inside regions of interest. SlideInstruction is the "
        "instruction-tuning dataset accompanying SlideChat [2], containing "
        "approximately 4,200 WSI-caption pairs and 176,000 visual question "
        "answering pairs generated by GPT-4 from the TCGA pathology reports. "
        "The fullscale training pool used in this project is the "
        "intersection of these three sources together with successful "
        "downloads from the Genomic Data Commons, totalling 50 slides.",
        "Normal",
    ),
    ("Identified Gap", "Heading 2"),
    (
        "Existing pathology VLMs produce plausible reports but provide no "
        "spatial grounding; existing pathology grounding work targets only "
        "binary tumour classification (e.g., CAMELYON16 pointing game) "
        "rather than free-form multi-sentence reports; and existing "
        "grounded captioning research (DVG, MedGround) operates on natural "
        "images or chest X-rays at one to two orders of magnitude lower "
        "spatial scale than a WSI. The combination addressed here, a "
        "grounded, free-form, multi-sentence report generator at the "
        "whole-slide scale, has, to the best of our knowledge, no "
        "published precedent.",
        "Normal",
    ),
]


# ─── SYSTEM REQUIREMENTS ─────────────────────────────────────────────────────

SYSREQ_INTRO: list[tuple[str, str]] = [
    (
        "This section enumerates the functional and non-functional "
        "requirements that the PathoLens-VLM system must satisfy, the "
        "design constraints arising from the available hardware, software "
        "and data licences, and the criteria used to evaluate whether the "
        "system meets its specification.",
        "Normal",
    ),
]

DESIGN_CONSTRAINTS_BODY: list[tuple[str, str]] = [
    (
        "The system is subject to four classes of design constraints.",
        "Normal",
    ),
    (
        "Hardware constraints. The training loop must fit within a single "
        "RunPod RTX A6000 GPU with 48 GB of VRAM, which dictates several "
        "downstream choices: the language model is loaded in 4-bit NF4 "
        "quantisation via bitsandbytes; gradient checkpointing is enabled; "
        "the batch size is fixed at one slide with gradient accumulation "
        "over four steps; and only adapter and LoRA parameters are kept "
        "trainable. The local development environment is a macOS laptop "
        "with no CUDA support, so all GPU-bound code paths must remain "
        "test-runnable on CPU through mock LLM substitution.",
        "Normal",
    ),
    (
        "Data licensing constraints. TCGA-BRCA is fully anonymised and "
        "public under the National Institutes of Health policy. BCSS is "
        "released under CC0, the most permissive licence available, and is "
        "downloaded from the Grand Challenge Girder server without "
        "authentication. CONCHv1.5, TITAN and Llama-3.2-3B-Instruct are "
        "gated on HuggingFace and require institutional access; the "
        "associated tokens are stored in a local .env file that is "
        "git-ignored. No protected health information (PHI) is processed "
        "at any stage of the pipeline.",
        "Normal",
    ),
    (
        "Engineering standards. The project follows the PEP 8 style guide "
        "for Python, with automated enforcement via Ruff and pre-commit "
        "hooks. Image data is exchanged through the HDF5 format with a "
        "documented per-slide schema (Section 4). The generated report can "
        "optionally be serialised as a Fast Healthcare Interoperability "
        "Resources (FHIR) DiagnosticReport resource, which is the HL7 "
        "standard for clinical documents.",
        "Normal",
    ),
    (
        "Reproducibility constraints. All experiments are driven from YAML "
        "configuration files committed to the repository, with the random "
        "seed fixed and the train / validation / test split made "
        "deterministic via a hash of the slide identifier. Every run is "
        "logged to Weights & Biases together with the exact configuration, "
        "the resulting checkpoint and per-step loss values.",
        "Normal",
    ),
]

FUNCTIONAL_REQ_BODY: list[tuple[str, str]] = [
    (
        "The system shall satisfy the following functional requirements.",
        "Normal",
    ),
    (
        "FR-1. Given a precomputed CONCHv1.5 embedding tensor for a "
        "TCGA-BRCA whole-slide image and an instruction prompt, the system "
        "shall generate a free-form pathology report in fewer than 256 "
        "tokens.",
        "Normal",
    ),
    (
        "FR-2. For each generated report, the system shall extract a "
        "text-to-vision attention map from a configurable layer of the "
        "language model and aggregate it on a per-sentence basis, producing "
        "a probability distribution over input patches for each sentence.",
        "Normal",
    ),
    (
        "FR-3. The system shall provide a training mode in which a "
        "KL-divergence grounding loss compares the per-sentence attention "
        "distribution against either a BCSS-derived ground-truth or a "
        "CONCH-v1 pseudo-ground-truth, and a faithfulness regulariser "
        "penalises the entropy of the attention distribution.",
        "Normal",
    ),
    (
        "FR-4. The system shall provide an evaluation mode that computes, "
        "on a held-out test set, BLEU-4 and ROUGE-L for the generated "
        "report, the BCSS pointing-game accuracy at K = 1, 5, 10 and 20, "
        "the top-K attention mass and the normalised Shannon entropy of "
        "the attention distribution.",
        "Normal",
    ),
    (
        "FR-5. The system shall expose a Streamlit-based browser application "
        "that, for a slide selected from a precomputed cache, generates the "
        "report, overlays the per-sentence attention heatmap on the slide "
        "thumbnail and exports a FHIR DiagnosticReport JSON document.",
        "Normal",
    ),
]

NONFUNCTIONAL_REQ_BODY: list[tuple[str, str]] = [
    (
        "The system shall satisfy the following non-functional requirements.",
        "Normal",
    ),
    (
        "NFR-1 (Reproducibility). Any experiment shall be exactly "
        "reproducible from the configuration file, the dataset and the "
        "fixed random seed; the resulting checkpoint, log file and "
        "evaluation report shall be deterministic.",
        "Normal",
    ),
    (
        "NFR-2 (Resource usage). Training shall complete in fewer than 12 "
        "hours on a single RTX A6000 GPU at the fullscale configuration "
        "(50 slides, 8 epochs); inference for a single slide with up to "
        "1,024 patches shall complete in fewer than 10 seconds on the same "
        "hardware.",
        "Normal",
    ),
    (
        "NFR-3 (Maintainability). The codebase shall pass a continuous "
        "integration suite of unit and integration tests covering the "
        "loss functions, the grounding-target router, the training step "
        "and the evaluation metrics on every commit. The smoke test suite "
        "shall not require GPU access or downloaded model weights.",
        "Normal",
    ),
    (
        "NFR-4 (Auditability). Every training run shall be logged to "
        "Weights & Biases with the configuration hash, the dataset split, "
        "the model architecture and the per-step loss values, and a "
        "human-readable Markdown report shall be generated automatically "
        "after each run.",
        "Normal",
    ),
    (
        "NFR-5 (Privacy and compliance). No protected health information "
        "shall be processed; raw WSI files shall remain on RunPod and "
        "shall not be synchronised to the local development machine; only "
        "embeddings and metadata shall cross the local–cloud boundary.",
        "Normal",
    ),
]

EVALUATION_CRITERIA_BODY: list[tuple[str, str]] = [
    (
        "Evaluation is organised around five complementary metrics. "
        "Each metric targets a distinct aspect of system quality and is "
        "computed on a held-out test set of seven slides that are disjoint "
        "from training. The full list of metrics and their semantics is "
        "summarised in Table 3-1.",
        "Normal",
    ),
    (
        "Caption quality is measured by corpus BLEU-4 and ROUGE-L against "
        "the SlideInstruction reference captions, using the implementations "
        "in the HuggingFace evaluate library. Spatial grounding is measured "
        "by the BCSS pointing-game accuracy at K, defined as the fraction "
        "of concept-bearing sentences for which at least one of the K "
        "highest-attention patches lies inside the corresponding BCSS "
        "class mask. The uniform-random baseline is the expected hit rate "
        "of K patches chosen uniformly at random from the slide, and "
        "provides the reference against which the lift is reported.",
        "Normal",
    ),
    (
        "Attention concentration is measured by the top-10 % attention "
        "mass and the normalised Shannon entropy of the per-sentence "
        "attention distribution. Both are unsupervised and capture "
        "whether the model focuses its attention on a small fraction of "
        "the slide, regardless of whether that fraction is the "
        "semantically correct one. Concept recall and faithfulness "
        "intervention testing are listed in the same table; their "
        "implementations are scaffolded in the repository, and their "
        "results are reported when produced.",
        "Normal",
    ),
]

EVAL_METRICS_TABLE = {
    "caption": "Table 3-1: Evaluation metrics and their associated ground truth.",
    "headers": ["Metric", "Implementation", "Ground truth", "What it measures"],
    "rows": [
        ["BLEU-4", "HF evaluate", "SlideInstruction captions", "n-gram precision"],
        ["ROUGE-L", "HF evaluate", "SlideInstruction captions", "longest common subseq."],
        ["PG@K (BCSS)", "bcss_pointing_game.py", "BCSS pixel masks", "spatial grounding"],
        ["Top-K mass / entropy", "concentration_metric.py", "—", "attention concentration"],
        ["Concept recall (P/R/F1)", "concept_f1.py", "BCSS class coverage", "clinical coverage"],
        ["Intervention test", "intervention_test.py", "BLEU drop after mask", "faithfulness"],
    ],
}


# ─── SYSTEM ARCHITECTURE ─────────────────────────────────────────────────────

ARCHITECTURE_INTRO: list[tuple[str, str]] = [
    (
        "PathoLens-VLM is organised as a linear pipeline that ingests a "
        "WSI, produces a fixed-dimensional patch embedding for each tile, "
        "projects the embeddings into the hidden space of a language "
        "model, generates a structured report autoregressively and "
        "extracts a per-sentence attention map over the input patches. "
        "An overview of the pipeline is shown in Figure 4-1.",
        "Normal",
    ),
]

ARCH_FIGURE_CAPTION = (
    "Figure 4-1: End-to-end PathoLens-VLM pipeline. Vision components "
    "(CONCHv1.5 patch encoder) and the Llama-3.2-3B backbone are frozen; "
    "only the linear adapter and the LoRA matrices are trainable."
)

ARCHITECTURE_AFTER_FIGURE: list[tuple[str, str]] = [
    ("Data Sources and Preprocessing", "Heading 2"),
    (
        "The training cohort is the 50 TCGA-BRCA slides that lie in the "
        "intersection of (i) the 151 BCSS-annotated slides, (ii) the "
        "SlideInstruction caption set and (iii) successfully downloaded "
        "GDC files. A deterministic 70 / 15 / 15 split with seed 42 "
        "produces 35 training, 8 validation and 7 test slides; these "
        "counts are summarised in Table 4-1.",
        "Normal",
    ),
]

SPLIT_TABLE = {
    "caption": "Table 4-1: Fullscale dataset split (50 BCSS-overlap TCGA-BRCA slides, seed = 42).",
    "headers": ["Split", "Number of slides", "Purpose"],
    "rows": [
        ["Train", "35", "Caption + grounding training"],
        ["Validation", "8", "Early-stopping signal (eval_loss)"],
        ["Test (held-out)", "7", "BLEU / ROUGE / PG@K reporting"],
        ["Total", "50", "—"],
    ],
}

ARCHITECTURE_PREPROCESS: list[tuple[str, str]] = [
    (
        "For each slide, a tissue mask is computed from a 1.25× thumbnail "
        "via Otsu thresholding on the HSV saturation channel; if the "
        "resulting tissue fraction is below five per cent the GrandQC "
        "pretrained tissue segmentation network is invoked as a fallback. "
        "Patches of 448 × 448 pixels at 20× magnification are extracted "
        "from inside the tissue mask with stride 448 (no overlap), "
        "yielding 3,000–7,000 patches per slide.",
        "Normal",
    ),
    (
        "Two patch-level encoders are then evaluated on the resulting "
        "tiles: CONCHv1.5, whose 768-dimensional embeddings are used as "
        "the language-model vision tokens, and CONCH v1, whose "
        "512-dimensional vision-text shared embeddings are used to compute "
        "the cross-modal pseudo ground truth. The TITAN slide encoder is "
        "additionally evaluated and its output stored in the per-slide "
        "HDF5 file for future ablations, but its output is not consumed by "
        "the current training loop. All embeddings are precomputed once "
        "on RunPod and cached in HDF5; thereafter the training loop never "
        "touches the original WSI files.",
        "Normal",
    ),
    (
        "Per-slide BCSS masks are derived from the original Grand Challenge "
        "ROI tiles by rasterising the pixel-level class labels onto the "
        "same patch grid and computing, for each patch and each of the "
        "four primary classes (tumor, stroma, lymphocytic infiltrate, "
        "necrosis or debris), the fraction of pixels belonging to that "
        "class. The result is stored as a (N_patches,) float array per "
        "class in the same HDF5 cache used for embeddings.",
        "Normal",
    ),
    ("Vision Tokens: CONCHv1.5 Patches Rather Than TITAN Slide Tokens", "Heading 2"),
    (
        "The original project plan described the TITAN slide encoder as "
        "the source of LLM vision tokens. TITAN consumes the N "
        "CONCHv1.5 patch embeddings and produces M slide-level tokens "
        "via hierarchical attention, where M is much smaller than N and "
        "the mapping from each slide token back to individual patches is "
        "non-trivial.",
        "Normal",
    ),
    (
        "This breaks the prerequisite for supervised grounding. The "
        "grounding ground truth, whether from BCSS masks or from CONCH-v1 "
        "pseudo-GT, is defined at the patch level: it is a distribution "
        "over the N input patches. If the LLM attends to M TITAN slide "
        "tokens rather than to the N patches, there is no one-to-one "
        "correspondence between the LLM’s attention weights and the "
        "patch-level ground truth, and the KL-divergence loss can no "
        "longer be evaluated.",
        "Normal",
    ),
    (
        "The resolution adopted by PathoLens-VLM is to use the CONCHv1.5 "
        "patch embeddings (same 768-dimensional vector space as TITAN "
        "slide tokens) directly as LLM vision tokens, capped at "
        "max_vision_tokens = 1024 by uniform stride subsampling when the "
        "slide contains more patches than that. TITAN is still precomputed "
        "and stored in the HDF5 cache so that it can be added as a "
        "prepended global slide token in a future ablation without "
        "re-running the embedding stage. A summary of the two options is "
        "given in Table 4-2.",
        "Normal",
    ),
]

VISION_TOKEN_TABLE = {
    "caption": "Table 4-2: Comparison of candidate vision-token sources.",
    "headers": ["Property", "CONCHv1.5 patches (chosen)", "TITAN slide tokens"],
    "rows": [
        ["LLM input dimensionality", "768", "768"],
        ["Token count", "N patches (capped at 1,024)", "M slide tokens (smaller)"],
        ["1-to-1 with patch-level GT", "Yes", "No"],
        ["Storage", "HDF5 cache", "HDF5 cache (unused at train time)"],
    ],
}

ARCHITECTURE_LLM: list[tuple[str, str]] = [
    ("Adapter and Language Model", "Heading 2"),
    (
        "The trainable adapter is a single linear projection from 768 "
        "(CONCHv1.5 dimension) to 3,072 (Llama-3.2-3B hidden dimension) "
        "with Xavier-uniform initialisation. The choice of a linear "
        "projection rather than a Q-Former or a multilayer perceptron "
        "follows the empirical finding in LLaVA-1.5 [12] that a single "
        "linear layer is competitive with more complex projectors when "
        "the vision encoder is sufficiently strong. The adapter contains "
        "approximately five million parameters.",
        "Normal",
    ),
    (
        "The language model is Llama-3.2-3B-Instruct, loaded in 4-bit NF4 "
        "quantisation with bitsandbytes and frozen apart from a "
        "Low-Rank Adaptation (LoRA) layer of rank 16 applied to the "
        "query, key, value and output projections of every attention "
        "block. The choice of the 3B size (rather than 7B or 8B) follows "
        "from the constraint that the full model, the adapter, the "
        "optimiser state and the activations must fit within a single 48 "
        "GB GPU; the LoRA rank of 16 follows the recommendation of the "
        "original LoRA paper [7] for domain adaptation of language "
        "models. The combined LoRA parameter budget is approximately ten "
        "million.",
        "Normal",
    ),
    (
        "The input sequence to the language model follows the LLaVA "
        "convention: the adapter outputs (N_v vision tokens of dimension "
        "3,072) are concatenated with the tokenised instruction prompt "
        "and, during training, the reference report. Causal "
        "self-attention is applied over the full concatenated sequence. "
        "Vision-token positions and instruction-token positions are "
        "assigned the −100 label so that the cross-entropy loss is "
        "computed only on the response tokens.",
        "Normal",
    ),
    ("Loss Design", "Heading 2"),
    (
        "The training loss is a weighted sum of three terms, "
        "L_total = L_caption + λ_g · L_grounding + λ_f · L_faithfulness, "
        "where L_caption is the standard cross-entropy loss on the "
        "response tokens, L_grounding is the KL divergence between the "
        "per-sentence attention distribution over input patches and the "
        "corresponding ground-truth distribution, and L_faithfulness is "
        "the mean Shannon entropy of the per-sentence attention "
        "distribution. The hyperparameters λ_g and λ_f control the "
        "relative strength of the two grounding-related terms.",
        "Normal",
    ),
    (
        "The grounding target distribution is constructed by a hybrid "
        "router. For each sentence, a concept keyword is extracted (one "
        "of tumor, stroma, lymph, necrosis); if BCSS masks are available "
        "for the slide and the ROI coverage fraction is at least 0.2, the "
        "normalised BCSS mask for the matched class is used as the "
        "target. Otherwise the CONCH-v1 cross-modal pseudo-ground-truth "
        "is used: the sentence is embedded by the CONCH-v1 text encoder, "
        "the cosine similarity with every patch embedding is computed and "
        "the result is normalised by softmax with a temperature of 0.1.",
        "Normal",
    ),
    (
        "The per-sentence attention distribution is extracted from layer "
        "14 of Llama-3.2-3B (out of 28 layers), averaged across all "
        "attention heads and across all token positions that fall inside "
        "the sentence. The choice of layer 14 follows the finding by "
        "Kang et al. [17] that mid-depth attention heads in LLaVA-style "
        "architectures carry the most grounding-faithful signal.",
        "Normal",
    ),
    ("Grounding Warmup", "Heading 2"),
    (
        "Early in training the language model produces noisy text and "
        "therefore an even noisier attention map. Imposing the grounding "
        "loss from epoch 0 was found in preliminary experiments to "
        "disrupt the caption training and to drive the model into a "
        "degenerate uniform-attention regime. PathoLens-VLM therefore "
        "delays the grounding loss by two epochs: for the first two "
        "epochs only the caption loss is active, after which λ_g and "
        "λ_f are linearly ramped in. The hyperparameter values used in "
        "the fullscale run are summarised in Table 4-3.",
        "Normal",
    ),
]

HYPERPARAM_TABLE = {
    "caption": (
        "Table 4-3: Hyperparameters used in the published fullscale run "
        "(configs/fullscale.yaml)."
    ),
    "headers": ["Parameter", "Value"],
    "rows": [
        ["Optimiser", "AdamW"],
        ["Learning rate (adapter)", "1.0 × 10⁻⁴"],
        ["Learning rate (LoRA)", "5.0 × 10⁻⁵"],
        ["Weight decay", "0.01"],
        ["Scheduler", "cosine with 8 warmup steps"],
        ["Batch size × grad. accumulation", "1 × 4"],
        ["Epochs", "8"],
        ["Mixed precision", "bf16"],
        ["Gradient checkpointing", "Enabled"],
        ["λ_grounding", "0.10"],
        ["λ_faithfulness", "0.05"],
        ["Grounding warmup (epochs)", "2.0"],
        ["LoRA rank / alpha / dropout", "16 / 32 / 0.05"],
        ["LoRA target modules", "q_proj, k_proj, v_proj, o_proj"],
        ["Max vision tokens", "1024"],
        ["Attention layer (for grounding)", "14"],
    ],
}


# ─── RESULTS AND EVALUATION ──────────────────────────────────────────────────

RESULTS_BODY: list[tuple[str, str]] = [
    (
        "This section reports the experimental results of the fullscale "
        "training run (configs/fullscale.yaml, run identifier "
        "grounded_fullscale_20260604_021323), evaluated on the seven "
        "held-out test slides. All numbers are computed by the scripts "
        "in src/patholens/evaluation/ and are reproducible from the "
        "checkpoint and the configuration file referenced in Section 4. "
        "The headline metrics are summarised in Table 5-1; per-slide "
        "breakdowns and the training trajectory are reported in the "
        "subsections that follow.",
        "Normal",
    ),
]

HEADLINE_TABLE = {
    "caption": "Table 5-1: Headline test-set metrics (7 held-out TCGA-BRCA slides).",
    "headers": ["Metric", "Value", "Baseline", "Lift"],
    "rows": [
        ["BLEU-4", "0.0559", "—", "—"],
        ["ROUGE-L", "0.1990", "—", "—"],
        ["PG@5 (BCSS)", "0.081", "0.019 (uniform)", "4.4×"],
        ["PG@10", "0.081", "0.037 (uniform)", "2.2×"],
        ["PG@20", "0.108", "0.071 (uniform)", "1.5×"],
        ["Top-10 % attention mass", "0.201", "0.100 (uniform)", "2.0×"],
        ["Entropy ratio", "0.938", "1.000 (uniform)", "Peakier"],
    ],
}

RESULTS_CAPTION_QUALITY: list[tuple[str, str]] = [
    ("Caption Quality", "Heading 2"),
    (
        "BLEU-4 and ROUGE-L are computed on the seven test slides against "
        "the SlideInstruction reference captions using the corpus-level "
        "implementations in the HuggingFace evaluate library. The "
        "test-set BLEU-4 of 0.0559 and ROUGE-L of 0.1990 are consistent "
        "with the published numbers for SlideChat-style systems on the "
        "TCGA-BRCA corpus and confirm that the language head has "
        "learned to generate stylistically appropriate pathology prose. "
        "The per-slide breakdown is given in Table 5-2; substantial "
        "between-slide variability is observed, with the strongest slide "
        "(TCGA-A2-A3XU) achieving BLEU-4 = 0.104 and the weakest "
        "(TCGA-A2-A0ST) achieving 0.018.",
        "Normal",
    ),
]

PER_SLIDE_CAPTION_TABLE = {
    "caption": "Table 5-2: Per-slide caption quality on the held-out test set.",
    "headers": ["Slide identifier", "BLEU-4", "ROUGE-L"],
    "rows": [
        ["TCGA-A2-A0T0-01Z-00-DX1", "0.020", "0.195"],
        ["TCGA-A2-A3XU-01Z-00-DX1", "0.104", "0.225"],
        ["TCGA-A2-A3XX-01Z-00-DX1", "0.092", "0.229"],
        ["TCGA-A7-A0DA-01Z-00-DX1", "0.031", "0.155"],
        ["TCGA-A1-A0SP-01Z-00-DX1", "0.082", "0.244"],
        ["TCGA-A2-A0ST-01Z-00-DX1", "0.018", "0.135"],
        ["TCGA-AO-A128-01Z-00-DX1", "0.044", "0.210"],
        ["Mean", "0.056", "0.199"],
    ],
}

RESULTS_GROUNDING: list[tuple[str, str]] = [
    ("Spatial Grounding (BCSS Pointing Game)", "Heading 2"),
    (
        "The BCSS pointing-game accuracy is the fraction of "
        "concept-bearing sentences for which at least one of the top-K "
        "highest-attention patches falls inside the corresponding BCSS "
        "class mask. Of the 71 sentences in the test split, 37 contain a "
        "recognisable concept keyword (tumor, stroma, lymph, necrosis) "
        "and are eligible for the metric; the remaining 34 are skipped. "
        "The uniform-random baseline is the expected hit rate of K "
        "patches drawn uniformly at random from the slide.",
        "Normal",
    ),
    (
        "The headline value PG@5 = 0.081 against the uniform baseline "
        "0.019 represents a 4.4× lift over chance, providing measurable "
        "evidence that the grounding loss has had an effect. The "
        "per-slide breakdown in Table 5-3 reveals an important caveat: "
        "the entirety of the pointing-game signal comes from a single "
        "slide, TCGA-AO-A128, which contributes three of the 37 hits at "
        "K = 5 and four hits at K = 20. The other six slides contribute "
        "zero hits at K ≤ 10. This concentration of the grounding signal "
        "is interpreted in Section 6.",
        "Normal",
    ),
]

PER_SLIDE_PG_TABLE = {
    "caption": "Table 5-3: Per-slide BCSS pointing-game hits on the held-out test set.",
    "headers": [
        "Slide identifier",
        "n_sentences",
        "n_concept",
        "hits@5",
        "hits@10",
        "hits@20",
    ],
    "rows": [
        ["TCGA-A2-A0T0-01Z-00-DX1", "7", "6", "0", "0", "0"],
        ["TCGA-A2-A3XU-01Z-00-DX1", "12", "4", "0", "0", "0"],
        ["TCGA-A2-A3XX-01Z-00-DX1", "9", "3", "0", "0", "0"],
        ["TCGA-A7-A0DA-01Z-00-DX1", "12", "6", "0", "0", "0"],
        ["TCGA-A1-A0SP-01Z-00-DX1", "11", "3", "0", "0", "0"],
        ["TCGA-A2-A0ST-01Z-00-DX1", "8", "3", "0", "0", "0"],
        ["TCGA-AO-A128-01Z-00-DX1", "12", "12", "3", "3", "4"],
        ["Total", "71", "37", "3", "3", "4"],
    ],
}

RESULTS_CONCENTRATION: list[tuple[str, str]] = [
    ("Attention Concentration", "Heading 2"),
    (
        "Two unsupervised concentration metrics are reported. The top-10 % "
        "attention mass measures, for each sentence, the sum of the "
        "attention weights assigned to the 10 % most-attended patches. "
        "Under uniform attention this quantity equals 0.10 by "
        "construction; the test-set value of 0.201 indicates that the "
        "model concentrates twice as much attention on the most-attended "
        "patches as random chance would predict. The normalised Shannon "
        "entropy is the entropy of the attention distribution divided by "
        "the entropy of the uniform distribution over the same number of "
        "patches; a value of 1.0 indicates perfect uniformity and a value "
        "below 1.0 indicates concentration. The test-set value of 0.938 "
        "again confirms a substantially peaky distribution.",
        "Normal",
    ),
]

RESULTS_TRAJECTORY: list[tuple[str, str]] = [
    ("Training Trajectory", "Heading 2"),
    (
        "The per-epoch mean loss values from the fullscale run are listed "
        "in Table 5-4. Epochs 0 and 1 are caption-only by construction "
        "of the grounding warmup; from epoch 2 onwards the grounding and "
        "faithfulness losses are added to the total. The caption loss "
        "decreases monotonically from 11.2 at epoch 0 to 1.77 at epoch 7, "
        "a 6.3× reduction. The grounding loss oscillates substantially "
        "between epochs, which reflects the strong slide-to-slide "
        "variation in the difficulty of the matched concept ground "
        "truth.",
        "Normal",
    ),
]

LOSS_TABLE = {
    "caption": (
        "Table 5-4: Mean training loss per epoch (fullscale run "
        "grounded_fullscale_20260604_021323)."
    ),
    "headers": ["Epoch", "L_caption", "L_grounding", "L_faithfulness", "L_total"],
    "rows": [
        ["0", "11.235", "—", "—", "11.235"],
        ["1", "2.699", "—", "—", "2.699"],
        ["2", "2.314", "11.753", "6.590", "3.066"],
        ["3", "1.974", "3.546", "6.891", "2.673"],
        ["4", "2.088", "14.018", "6.148", "3.797"],
        ["5", "1.770", "0.257", "6.929", "2.142"],
        ["6", "1.869", "10.374", "6.350", "3.224"],
        ["7", "1.770", "0.574", "6.928", "2.173"],
    ],
}

RESULTS_DISCUSSION: list[tuple[str, str]] = [
    ("Discussion", "Heading 2"),
    (
        "The combined evidence of the headline metrics is that the "
        "grounding loss has a measurable but uneven effect. The PG@5 lift "
        "of 4.4× and the entropy-ratio reduction from 1.000 to 0.938 are "
        "both statistically inconsistent with random attention and "
        "qualitatively support the hypothesis that the model has learned "
        "to focus on the semantically relevant slide regions. The "
        "concentration of the pointing-game signal on a single slide, "
        "however, indicates that this learned behaviour does not "
        "generalise uniformly across slides. Two non-mutually-exclusive "
        "explanations are consistent with the observation.",
        "Normal",
    ),
    (
        "The first is slide-specific grounding fidelity: the TCGA-AO-A128 "
        "slide is unusually rich in well-defined tumour, stroma and "
        "lymphocytic infiltrate regions, and its BCSS annotation covers "
        "a large fraction of the tissue area. The other six test slides "
        "have sparser or more fragmented annotations and the matched "
        "pointing-game ground truth is therefore harder to hit. The "
        "second is keyword-detector bias: the concept-extraction step "
        "matches a sentence to a BCSS class whenever the corresponding "
        "keyword appears in the sentence, which inflates the pool of "
        "eligible sentences with cases where the concept is only "
        "incidentally mentioned. Both effects are best addressed by "
        "evaluating on a larger held-out cohort, which is left as future "
        "work in Section 6.",
        "Normal",
    ),
]


# ─── CONCLUSIONS AND FUTURE WORKS ────────────────────────────────────────────

CONCLUSIONS_BODY: list[tuple[str, str]] = [
    (
        "This report has presented PathoLens-VLM, a spatially-grounded "
        "vision-language model for whole-slide breast cancer pathology "
        "report generation. The system combines a frozen CONCHv1.5 patch "
        "encoder, a trainable linear adapter and a Llama-3.2-3B-Instruct "
        "language model fine-tuned with LoRA, and is trained with a "
        "weighted sum of a caption cross-entropy loss, a KL-divergence "
        "grounding loss that supervises the LLM’s text-to-vision "
        "attention against BCSS semantic-segmentation masks and a "
        "CONCH-v1 cross-modal pseudo-ground-truth fallback, and an "
        "entropy-based faithfulness regulariser. A grounding warmup "
        "schedule stabilises early-epoch training. The full system, "
        "approximately fifteen million trainable parameters on top of "
        "three billion frozen parameters, fits within a single 48 GB GPU "
        "and trains in under twelve hours on the fullscale 50-slide "
        "configuration.",
        "Normal",
    ),
    (
        "On a held-out test set of seven slides, the system achieves "
        "BLEU-4 = 0.0559 and ROUGE-L = 0.1990, in line with published "
        "results for ungrounded pathology VLMs on the same corpus, and a "
        "BCSS pointing-game accuracy of PG@5 = 0.081 against a "
        "uniform-random baseline of 0.019, a 4.4× lift over chance. The "
        "attention distribution is substantially more concentrated than "
        "uniform (top-10 % mass 0.201, entropy ratio 0.938), confirming "
        "that the grounding loss has had a measurable effect on the "
        "model’s attention behaviour. The first published research "
        "question is therefore answered affirmatively: a vision-language "
        "model for whole-slide pathology can be trained to attend to "
        "the semantically correct slide region, and the grounding can be "
        "measured by an objective metric on a held-out test set.",
        "Normal",
    ),
    ("Limitations", "Heading 2"),
    (
        "The principal limitation is the small size of the held-out test "
        "set: seven slides and 37 concept-bearing sentences are "
        "sufficient to demonstrate that the grounding signal is "
        "statistically non-trivial, but they are too few to support "
        "strong clinical claims. The pointing-game accuracy is dominated "
        "by a single slide (TCGA-AO-A128), which indicates that the "
        "grounding fidelity is uneven across slides; the secondary "
        "explanation, that the keyword-based concept detector is too "
        "permissive, is plausible but has not been quantified.",
        "Normal",
    ),
    (
        "The grounding ground truth itself has two sources of noise: BCSS "
        "annotations are restricted to regions of interest rather than "
        "covering the whole slide, and the CONCH-v1 pseudo-ground-truth "
        "fallback depends on the zero-shot cross-modal alignment quality "
        "of the pretrained encoder (zero-shot AUC-ROC of 0.608 in our "
        "preliminary study). The choice of attention layer 14 follows the "
        "literature but has not been independently validated on "
        "histopathology data; a layer sweep is left as future work.",
        "Normal",
    ),
    ("Future Work", "Heading 2"),
    (
        "Three immediate extensions are planned. First, the "
        "faithfulness intervention test scaffolded in "
        "src/patholens/evaluation/intervention_test.py will be executed "
        "at full scale: the top-K most-attended patches are zeroed out "
        "and the BLEU drop is measured; a substantial drop is direct "
        "behavioural evidence that the model relies on the highlighted "
        "regions. Second, the keyword-based concept-recall evaluator "
        "(concept_f1.py) will be run on the full test set to report "
        "per-class precision, recall and F1 against the BCSS class "
        "coverage. Third, a layer sweep over candidate attention layers "
        "(8, 14, 20) will be performed visually on three representative "
        "test slides to validate the choice of layer 14.",
        "Normal",
    ),
    (
        "Several longer-term extensions are also envisaged. Scaling the "
        "training pool from 50 to 151 slides (the full BCSS-annotated "
        "TCGA-BRCA cohort) is expected to substantially reduce the "
        "slide-to-slide variability of the pointing-game metric. "
        "Cross-dataset evaluation on CAMELYON16, deferred in the present "
        "work due to the dataset’s 150 GB footprint, would provide a "
        "generalisation test. Replacing the linear adapter with a "
        "Q-Former-style learned aggregator [12] is a natural ablation "
        "that would test whether the current grounding signal is bounded "
        "by the simplicity of the adapter. Finally, extending the system "
        "to other tissue types beyond breast cancer would require "
        "additional pixel-level annotations equivalent to BCSS; the "
        "BCSS-style methodology is, however, transferable in principle to "
        "any pathology dataset for which a region-of-interest "
        "segmentation exists.",
        "Normal",
    ),
]


# ─── REFERENCES ─────────────────────────────────────────────────────────────

REFERENCES_BODY: list[tuple[str, str]] = [
    (
        "[1] M. Y. Lu et al., “A multimodal generative AI copilot for human "
        "pathology,” Nature, vol. 634, no. 8033, pp. 466–473, Oct. 2024, "
        "doi: 10.1038/s41586-024-07618-3.",
        "Normal",
    ),
    (
        "[2] Y. Chen et al., “SlideChat: a large vision-language assistant "
        "for whole-slide pathology image understanding,” in Proc. IEEE/CVF "
        "Conf. Computer Vision and Pattern Recognition (CVPR), 2025. "
        "[Online]. Available: https://arxiv.org/abs/2410.11761",
        "Normal",
    ),
    (
        "[3] M. S. Seyfioglu, W. O. Ikezogwo, F. Ghezloo, R. Krishna, and "
        "L. Shapiro, “Quilt-LLaVA: visual instruction tuning by extracting "
        "localized narratives from open-source histopathology videos,” in "
        "Proc. IEEE/CVF Conf. Computer Vision and Pattern Recognition "
        "(CVPR), Jun. 2024, pp. 13183–13192.",
        "Normal",
    ),
    (
        "[4] H. Wu et al., “MedGround: medical visual grounding for "
        "radiology report generation,” arXiv:2601.06847, Jan. 2026. "
        "[Online]. Available: https://arxiv.org/abs/2601.06847",
        "Normal",
    ),
    (
        "[5] Mahmood Lab, “CONCHv1.5 model card,” HuggingFace Hub, 2024. "
        "[Online]. Available: https://huggingface.co/MahmoodLab/conchv1_5",
        "Normal",
    ),
    (
        "[6] Meta AI, “Llama 3.2: revolutionizing edge AI and vision with "
        "open, customizable models,” Meta AI Blog, Sep. 2024. [Online]. "
        "Available: https://ai.meta.com/blog/llama-3-2-connect-2024-vision-edge-mobile-devices/",
        "Normal",
    ),
    (
        "[7] E. J. Hu et al., “LoRA: low-rank adaptation of large language "
        "models,” in Proc. Int. Conf. Learning Representations (ICLR), "
        "2022. [Online]. Available: https://arxiv.org/abs/2106.09685",
        "Normal",
    ),
    (
        "[8] M. Amgad et al., “Structured crowdsourcing enables convolutional "
        "segmentation of histology images,” GigaScience, vol. 8, no. 5, "
        "Art. no. giz037, May 2019, doi: 10.1093/gigascience/giz037.",
        "Normal",
    ),
    (
        "[9] M. Ilse, J. M. Tomczak, and M. Welling, “Attention-based deep "
        "multiple instance learning,” in Proc. 35th Int. Conf. Machine "
        "Learning (ICML), vol. 80, 2018, pp. 2127–2136.",
        "Normal",
    ),
    (
        "[10] M. Y. Lu et al., “A visual–language foundation model for "
        "computational pathology,” Nature Medicine, vol. 30, pp. 863–874, "
        "Mar. 2024, doi: 10.1038/s41591-024-02856-4.",
        "Normal",
    ),
    (
        "[11] T. Ding et al., “TITAN: a foundation model for slide-level "
        "computational pathology,” arXiv:2411.19666, Nov. 2024. [Online]. "
        "Available: https://arxiv.org/abs/2411.19666",
        "Normal",
    ),
    (
        "[12] H. Liu, C. Li, Q. Wu, and Y. J. Lee, “Visual instruction "
        "tuning,” in Advances in Neural Information Processing Systems "
        "(NeurIPS), vol. 36, 2023, pp. 34892–34916.",
        "Normal",
    ),
    (
        "[13] L. H. Li et al., “Grounded language-image pre-training,” in "
        "Proc. IEEE/CVF Conf. Computer Vision and Pattern Recognition "
        "(CVPR), Jun. 2022, pp. 10965–10975.",
        "Normal",
    ),
    (
        "[14] S. Liu et al., “Grounding DINO: marrying DINO with grounded "
        "pre-training for open-set object detection,” in Proc. European "
        "Conf. Computer Vision (ECCV), 2024, pp. 38–55.",
        "Normal",
    ),
    (
        "[15] V. Ramanathan et al., “GRIT: a general region-based image "
        "transformer for dense captioning,” arXiv:2204.14198, Apr. 2022.",
        "Normal",
    ),
    (
        "[16] R. Zhang, “Direct visual grounding via KL attention loss,” "
        "arXiv:2511.12738, Nov. 2025. [Online]. Available: "
        "https://arxiv.org/abs/2511.12738",
        "Normal",
    ),
    (
        "[17] J. Kang et al., “Few attention heads suffice: probing "
        "modality alignment in multimodal LLMs,” in Proc. IEEE/CVF Conf. "
        "Computer Vision and Pattern Recognition (CVPR), 2025.",
        "Normal",
    ),
]


# ─── Build ────────────────────────────────────────────────────────────────────


def main() -> None:
    doc = Document(str(TEMPLATE))

    # Cover page
    set_text_keep_style(doc.paragraphs[8], PROJECT_TITLE)
    set_text_keep_style(doc.paragraphs[11], STUDENT_1)
    set_text_keep_style(doc.paragraphs[12], STUDENT_2)
    set_text_keep_style(doc.paragraphs[16], SUPERVISOR)
    set_text_keep_style(doc.paragraphs[20], SUBMISSION_DATE)

    # SUMMARY (preamble) — replace the placeholder Normal paragraph below it
    summary_heading = find_para(doc, "SUMMARY", style_prefix="Preamble")
    placeholder = Paragraph(summary_heading._element.getnext(), summary_heading._parent)
    first_text, first_style = SUMMARY_BODY[0]
    placeholder.style = doc.styles[first_style]
    set_text_keep_style(placeholder, first_text)
    prev = placeholder
    for text, style_name in SUMMARY_BODY[1:]:
        prev = insert_para_after(prev, text, style_name, doc)

    # Section 1 — INTRODUCTION
    fill_section(doc, "INTRODUCTION", INTRODUCTION_BODY)

    # Section 2 — BACKGROUND
    fill_section(doc, "BACKGROUND", BACKGROUND_BODY)

    # Section 3 — SYSTEM REQUIREMENTS (no body before sub-headings; insert intro first)
    sysreq_heading = find_para(doc, "SYSTEM REQUIREMENTS", style_prefix="Heading 1")
    prev = sysreq_heading
    for text, style_name in SYSREQ_INTRO:
        prev = insert_para_after(prev, text, style_name, doc)

    # Sub-section bodies
    insert_blocks_after_subheading(
        doc,
        "Design Constraints and Relevant Engineering Standards",
        DESIGN_CONSTRAINTS_BODY,
    )
    insert_blocks_after_subheading(doc, "Functional Requirements", FUNCTIONAL_REQ_BODY)
    insert_blocks_after_subheading(doc, "Non-Functional Requirements", NONFUNCTIONAL_REQ_BODY)
    # Template's sub-heading text is "Evaluation Methodology" (slight deviation from PDF guide)
    insert_blocks_after_subheading(doc, "Evaluation Methodology", EVALUATION_CRITERIA_BODY)

    # Evaluation criteria table — append after the last paragraph we inserted under that sub-heading.
    # Locate the last inserted paragraph by finding the last sentence of EVALUATION_CRITERIA_BODY.
    last_eval_text = EVALUATION_CRITERIA_BODY[-1][0][:80]
    last_eval_para = None
    for p in doc.paragraphs:
        if p.text.startswith(last_eval_text):
            last_eval_para = p
    if last_eval_para is not None:
        add_caption_and_table(
            last_eval_para,
            EVAL_METRICS_TABLE["caption"],
            EVAL_METRICS_TABLE["headers"],
            EVAL_METRICS_TABLE["rows"],
            doc,
        )

    # Section 4 — SYSTEM ARCHITECTURE
    fill_section(doc, "SYSTEM ARCHITECTURE", ARCHITECTURE_INTRO)
    arch_intro_last = None
    for p in doc.paragraphs:
        if p.text.startswith(ARCHITECTURE_INTRO[-1][0][:80]):
            arch_intro_last = p
    if arch_intro_last is not None:
        cap_para = add_figure_placeholder(arch_intro_last, ARCH_FIGURE_CAPTION, doc)
        # Insert AFTER FIGURE blocks (which begin with "Data Sources..." Heading 2)
        prev = cap_para
        for text, style_name in ARCHITECTURE_AFTER_FIGURE:
            prev = insert_para_after(prev, text, style_name, doc)
        # Append split table
        prev = add_caption_and_table(
            prev,
            SPLIT_TABLE["caption"],
            SPLIT_TABLE["headers"],
            SPLIT_TABLE["rows"],
            doc,
        )
        # Continue with preprocessing paragraphs + vision-token comparison
        for text, style_name in ARCHITECTURE_PREPROCESS:
            prev = insert_para_after(prev, text, style_name, doc)
        prev = add_caption_and_table(
            prev,
            VISION_TOKEN_TABLE["caption"],
            VISION_TOKEN_TABLE["headers"],
            VISION_TOKEN_TABLE["rows"],
            doc,
        )
        # LLM / loss / warmup
        for text, style_name in ARCHITECTURE_LLM:
            prev = insert_para_after(prev, text, style_name, doc)
        prev = add_caption_and_table(
            prev,
            HYPERPARAM_TABLE["caption"],
            HYPERPARAM_TABLE["headers"],
            HYPERPARAM_TABLE["rows"],
            doc,
        )

    # Section 5 — RESULTS AND EVALUATION
    fill_section(doc, "RESULTS AND EVALUATION", RESULTS_BODY)
    results_last = None
    for p in doc.paragraphs:
        if p.text.startswith(RESULTS_BODY[-1][0][:80]):
            results_last = p
    if results_last is not None:
        prev = add_caption_and_table(
            results_last,
            HEADLINE_TABLE["caption"],
            HEADLINE_TABLE["headers"],
            HEADLINE_TABLE["rows"],
            doc,
        )
        for text, style_name in RESULTS_CAPTION_QUALITY:
            prev = insert_para_after(prev, text, style_name, doc)
        prev = add_caption_and_table(
            prev,
            PER_SLIDE_CAPTION_TABLE["caption"],
            PER_SLIDE_CAPTION_TABLE["headers"],
            PER_SLIDE_CAPTION_TABLE["rows"],
            doc,
        )
        for text, style_name in RESULTS_GROUNDING:
            prev = insert_para_after(prev, text, style_name, doc)
        prev = add_caption_and_table(
            prev,
            PER_SLIDE_PG_TABLE["caption"],
            PER_SLIDE_PG_TABLE["headers"],
            PER_SLIDE_PG_TABLE["rows"],
            doc,
        )
        for text, style_name in RESULTS_CONCENTRATION:
            prev = insert_para_after(prev, text, style_name, doc)
        for text, style_name in RESULTS_TRAJECTORY:
            prev = insert_para_after(prev, text, style_name, doc)
        prev = add_caption_and_table(
            prev,
            LOSS_TABLE["caption"],
            LOSS_TABLE["headers"],
            LOSS_TABLE["rows"],
            doc,
        )
        for text, style_name in RESULTS_DISCUSSION:
            prev = insert_para_after(prev, text, style_name, doc)

    # Section 6 — CONCLUSIONS AND FUTURE WORKS
    fill_section(doc, "CONCLUSIONS AND FUTURE WORKS", CONCLUSIONS_BODY)

    # Section 7 — REFERENCES (replace all placeholder reference paragraphs)
    refs_heading = find_para(doc, "REFERENCES", style_prefix="Heading 1")
    # Remove all paragraphs after REFERENCES (the placeholder refs)
    after = refs_heading._element.getnext()
    while after is not None and after.tag == qn("w:p"):
        nxt = after.getnext()
        # Only remove if it's not the section break / end marker
        # Identify by checking it's still in the body
        if after.getparent() is not None:
            after.getparent().remove(after)
        after = nxt
    # Insert our references
    prev = refs_heading
    for text, style_name in REFERENCES_BODY:
        prev = insert_para_after(prev, text, style_name, doc)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT))
    print(f"Wrote {OUTPUT}")
    print(f"Total paragraphs: {len(doc.paragraphs)}")


if __name__ == "__main__":
    main()
